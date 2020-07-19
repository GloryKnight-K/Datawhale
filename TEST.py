#实验现场日志
#需要的包 pip install docxtpl -i https://pypi.tuna.tsinghua.edu.cn/simple/ &&  pip install windrose

import pandas as pd
import numpy as np
from docx import document  #该模块可提供代码提示
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH #对齐
from docx.shared import Pt,Inches,Cm # 单位
from docx.oxml.ns import qn # 中文字体
from datetime import datetime
import matplotlib.pyplot as plt
from matplotlib.pyplot import MultipleLocator
from windrose import WindroseAxes
import matplotlib.cm as cm


#载入数据
#data=pd.read_htlm('接口网址')
Position='127.4E,54,7N'
DateT=datetime(2020,6,27,8)
AverageTemperature='28℃'
HighestTemperature='31℃'
TimeHT='14:30'
MinimumTemperature='15℃'
TimeMT='04:00'
AverageWindSpeed='4m/s'
MaximumWindSpeed='15m/s'
TimeMWS='16:00'
DayAccumulatedRainfall='2.5mm'
HourMaximumRainfall='1.1mm'
TimeHMR='13:00~14;00'
HourlyWindSpeed=[3,4,8,5,6,7,13,12,6,4,6,15,18]
x=range(0,25,2)

#风速折线图
plt.plot(x,HourlyWindSpeed,c='orange',marker="o")
plt.tick_params(axis='both',which='major',labelsize=14)
x_major_locator=MultipleLocator(2)
#把x轴的刻度间隔设置为2，并存在变量里
y_major_locator=MultipleLocator(5)
#把y轴的刻度间隔设置为5，并存在变量里
ax=plt.gca()
ax.xaxis.set_major_locator(x_major_locator)
ax.yaxis.set_major_locator(y_major_locator)
plt.xlim(-0.5,24)
plt.ylim(-1,20)
plt.legend()
plt.savefig(r'D:\windspeed.png', dpi=300,bbox_inches='tight')
plt.show()

#风向玫瑰图
plt.clf()
ws = np.random.random(500) * 6
wd = np.random.random(500) * 360
ax = WindroseAxes.from_ax()
ax.bar(wd, ws, normed=True, opening=0.8, edgecolor='white')
ax.set_legend()
plt.savefig(r'D:\rose.png', dpi=300,bbox_inches='tight')
plt.show()

#降水柱形图
Rainfall=[0,0,0,1,1,2,3,4,6,5,4,1,1]
x_major_locator=MultipleLocator(2)
#把x轴的刻度间隔设置为2，并存在变量里
y_major_locator=MultipleLocator(1)
#把y轴的刻度间隔设置为1，并存在变量里
ax=plt.gca()
ax.xaxis.set_major_locator(x_major_locator)
ax.yaxis.set_major_locator(y_major_locator)
plt.xlim(-0.5,25)
plt.ylim(0,6)
plt.bar(x,Rainfall)
plt.legend()
plt.savefig(r'D:\rainfall.png', dpi=300,bbox_inches='tight')
plt.show()

##气压折线图
AtmosphericPressure=[800,805,808,810,806,805,810,812,815,813,810,809,804]
plt.plot(x,AtmosphericPressure,c='purple',marker="o")
plt.tick_params(axis='both',which='major',labelsize=14)
x_major_locator=MultipleLocator(2)
#把x轴的刻度间隔设置为2，并存在变量里
y_major_locator=MultipleLocator(300)
#把y轴的刻度间隔设置为5，并存在变量里
ax=plt.gca()
ax.xaxis.set_major_locator(x_major_locator)
ax.yaxis.set_major_locator(y_major_locator)
plt.xlim(-0.5,25)
plt.ylim(0,1200)
plt.legend()
plt.savefig(r'D:\apressure.png', dpi=300,bbox_inches='tight')
plt.show()

doc=Document() # 新建文档
#doc=document.Document()   #为了代码提示，这句话在程序里是报错的
doc.styles['heading 1'].element.rPr.rFonts.set(qn('w:eastAsia'), 'Heiti SC Medium')
doc.styles['heading 1'].font.size=Pt(16)

gara1 = doc.add_paragraph() # 增加一个段落，这个段落是标题
gara1.alignment = WD_ALIGN_PARAGRAPH.CENTER # 对齐方式为居中，没有这句话默认左对齐
Topic= gara1.add_run('实验现场日志') # 给段落添加一个块，写上文字
Topic.font.name = 'Arial' # 设置英文字体
Topic.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体') # 设置中文字体
Topic.font.size = Pt(48) # 设置字号
Topic.font.bold = True # 设置加粗
gara1.space_after = Pt(5) # 断后距离5磅
gara1.space_before = Pt(5) # 断前距离5磅


table = doc.add_table(rows = 1, cols = 2)
table.rows[0].cells[0].text='试验位置：'+Position
table.rows[0].cells[1].text=str(DateT.year)+'年'+str(DateT.month)+'月'+str(DateT.day)+'日8:00'
for row in table.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(16)
                font.name = u'仿宋_GB2312'


heading1=doc.add_heading('一、天气概况',level=1)
gara2 = doc.add_paragraph()# 增加第二个段落，这个段落是正文
Summary= gara2.add_run(f'{DateT.month}月{DateT.day}日实验现场，当日平均气温{AverageTemperature}，其中最高气温{HighestTemperature}，出现在{TimeHT}；最低气温{MinimumTemperature}，出现在{TimeMT}。当日平均风速{AverageWindSpeed},最大风速{MaximumWindSpeed}，出现在{TimeMWS}。当日盛行风向西北风，详见当日风玫瑰图。当日24小时累计降水{DayAccumulatedRainfall}，小时最大降水为{HourMaximumRainfall}，出现在{TimeHMR}。\n')
Summary.font.name = '仿宋_GB2312'
Summary.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
Summary.font.size = Pt(16)
gara2.paragraph_format.first_line_indent = Inches(0.4) # 左缩进0.4英寸

heading2=doc.add_heading('二、地面天气实况曲线',level=1)

para3= doc.add_paragraph() # 增加一个自然段，这段是特此说明
Numbering1=para3.add_run('1. 风向风速')
Numbering1.font.name = '仿宋_GB2312'
Numbering1.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
Numbering1.font.size = Pt(16)
para3.paragraph_format.first_line_indent = Inches(0.4) # 左缩进0.4英寸 
para4= doc.add_paragraph() # 增加一个自然段，这段是特此说明
run=para4.add_run()
run.add_picture("D:/windspeed.png",width=Cm(8),height=Cm(6))
run.add_picture("D:/rose.png",width=Cm(6),height=Cm(6))
skip=doc.add_paragraph()
skip.add_run('\n\n\n')
para5= doc.add_paragraph() # 增加一个自然段，这段是特此说明
Numbering2=para5.add_run('2. 降水、气压')
Numbering2.font.name = '仿宋_GB2312'
Numbering2.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
Numbering2.font.size = Pt(16)
para5.paragraph_format.first_line_indent = Inches(0.4) # 左缩进0.4英寸 
para6= doc.add_paragraph() # 增加一个自然段，这段是特此说明
run=para6.add_run()
run.add_picture("D:/rainfall.png",width=Cm(7.5),height=Cm(5.6))
run.add_picture("D:/apressure.png",width=Cm(7.5),height=Cm(5.6))

doc.save('D:/实验现场日志.docx')